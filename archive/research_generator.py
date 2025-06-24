from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass, field
from enum import Enum
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

import docx2pdf
import win32com.client
import os
import datetime




class AlignmentType(Enum):
    """Text alignment options"""
    LEFT = WD_ALIGN_PARAGRAPH.LEFT
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    RIGHT = WD_ALIGN_PARAGRAPH.RIGHT
    

@dataclass
class FontStyle:
    """Font styling configuration"""
    name: str = 'Calibri'
    size: int = 10
    bold: bool = False
    italic: bool = False

@dataclass
class BorderStyle:
    """Border styling configuration"""
    size: str = "8"
    style: str = "single"
    color: str = "C0C0C0"

@dataclass
class TableStyle:
    """Table styling configuration"""
    font: FontStyle = field(default_factory=FontStyle)
    border: BorderStyle = field(default_factory=BorderStyle)
    alternate_shading: bool = True
    shade_color: str = "F2F2F2"
    row_height: float = 0.2
    bold_header: bool = False
    shade_header: bool = False


@dataclass 
class CompanyInfo:
    """Company information data structure"""
    name: str
    country: str
    analyst: str
    gics_sector: str
    gics_industry_group: str
    gics_industry: str
    gics_sub_industry: str
    update_date: str = field(default_factory=lambda: datetime.datetime.today().strftime('%d-%b-%Y'))

@dataclass
class InvestmentRecommendation:
    """Investment recommendation data structure"""
    buy_sell_rec: str
    idea_stage: str
    esg_rating: str
    theme: str
    last_price: float
    base_target: float
    bull_target: float
    bear_target: float

class DocumentStyler:
    """Handles all document styling operations"""
    
    def __init__(self, style_config: TableStyle = None):
        self.style_config = style_config or TableStyle()
    
    def set_cell_borders(self, cell: _Cell) -> None:
        """Apply borders to a table cell"""
        border_style = {
            "sz": self.style_config.border.size,
            "val": self.style_config.border.style,
            "color": self.style_config.border.color
        }
        
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Remove existing borders
        for border_tag in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(border_tag)
        
        # Add new borders
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ("start", "top", "end", "bottom"):
            element = OxmlElement(f"w:{edge}")
            for key, val in border_style.items():
                element.set(qn(f"w:{key}"), val)
            tcBorders.append(element)
        tcPr.append(tcBorders)
    
    def set_cell_shading(self, cell: _Cell, fill_color: str) -> None:
        """Apply shading to a table cell"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        tcPr.append(shd)
    
    def style_cell_text(self, cell: _Cell, text: str, 
                       alignment: AlignmentType = AlignmentType.LEFT,
                       font_override: FontStyle = None) -> None:
        """Apply text styling to a cell"""
        font_style = font_override or self.style_config.font
        
        p = cell.paragraphs[0]
        p.alignment = alignment.value
        
        # Clear existing runs
        p.clear()
        
        run = p.add_run(text)
        run.font.name = font_style.name
        run.font.size = Pt(font_style.size)
        run.bold = font_style.bold
        run.italic = font_style.italic

    def style_table(self, table: Table, table_type: str = "default") -> None:
        """Apply comprehensive styling to a table"""
        for row_idx, row in enumerate(table.rows):
            # Set row height (except for content rows in thesis tables)
            if table_type == "thesis" and row_idx == 1:  # Content row in thesis table
                # Don't set fixed height for thesis content - let it expand
                pass
            else:
                row.height = Inches(self.style_config.row_height)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            # Determine shading based on table type and row
            should_shade = False
            
            if table_type == "key_value":
                # For key-value tables, shade every other row starting with first
                should_shade = (self.style_config.alternate_shading and row_idx % 2 == 0)
            elif table_type == "thesis":
                # For thesis tables, shade the header row
                should_shade = (row_idx == 0 and self.style_config.shade_header)
            elif table_type == "data":
                # For data tables, shade alternate rows but not header
                should_shade = (self.style_config.alternate_shading and 
                              row_idx % 2 == 0 and row_idx > 0)
            
            for cell in row.cells:
                self.set_cell_borders(cell)
                if should_shade:
                    self.set_cell_shading(cell, self.style_config.shade_color)



class TableBuilder:
    """Builds different types of tables with consistent styling"""
    
    def __init__(self, styler: DocumentStyler):
        self.styler = styler
    
    def build_key_value_table(self, doc: Document, data: List[Tuple[str, str]]) -> Table:
        """Build a two-column key-value table"""
        table = doc.add_table(rows=len(data), cols=2)
        
        for row_idx, (key, value) in enumerate(data):
            # Key column
            key_cell = table.cell(row_idx, 0)
            self.styler.style_cell_text(key_cell, f"{key}:", AlignmentType.LEFT)
            
            # Value column (italic)
            value_cell = table.cell(row_idx, 1)
            value_font = FontStyle(italic=True)
            self.styler.style_cell_text(value_cell, value, AlignmentType.LEFT, value_font)
        
        self.styler.style_table(table, "key_value")
        self._set_column_widths(table, [Inches(3.75), Inches(3.75)])
        return table
    
    def build_data_table(self, doc: Document, data: List[List[str]], 
                        has_header: bool = True) -> Table:
        """Build a multi-column data table"""
        if not data:
            raise ValueError("Data cannot be empty")
        
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        
        for row_idx, row_data in enumerate(data):
            is_header = has_header and row_idx == 0
            
            for col_idx, cell_value in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                
                # Determine alignment
                if is_header:
                    if col_idx == 0:
                        alignment = AlignmentType.LEFT
                    else:
                        alignment = AlignmentType.CENTER
                    font_style = FontStyle(bold=True)
                elif col_idx == 0:
                    alignment = AlignmentType.LEFT
                    font_style = FontStyle()
                else:
                    alignment = AlignmentType.RIGHT
                    font_style = FontStyle()
                
                self.styler.style_cell_text(cell, str(cell_value), alignment, font_style)
        
        self.styler.style_table(table, "data")
        return table

    def build_single_column_table(self, doc: Document, data: List[str]) -> Table:
        """Build a single-column table (useful for thesis/description blocks)"""
        table = doc.add_table(rows=len(data), cols=1)
        
        for row_idx, text in enumerate(data):
            cell = table.cell(row_idx, 0)
            font_style = FontStyle(bold=(row_idx == 0))  # First row bold
            self.styler.style_cell_text(cell, text, AlignmentType.LEFT, font_style)
        
        self.styler.style_table(table, "thesis")
        return table
    
    def _set_column_widths(self, table: Table, widths: List[Inches]) -> None:
        """Set column widths for a table"""
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(widths):
                    cell.width = widths[col_idx]



class BloombergReportGenerator:
    """Main class for generating Bloomberg RMS NOTEs"""
    
    def __init__(self, style_config: TableStyle = None):
        self.styler = DocumentStyler(style_config)
        self.table_builder = TableBuilder(self.styler)
        self.doc = None
    
    def create_document(self, margins: float = 0.5) -> Document:
        """Create a new document with custom margins"""
        self.doc = Document()
        
        # Set margins
        section = self.doc.sections[0]
        section.top_margin = Inches(margins)
        section.bottom_margin = Inches(margins)
        section.left_margin = Inches(margins)
        section.right_margin = Inches(margins)

        # Add a header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = header_para.add_run("Bloomberg")
        run.font.name = 'Calibri (Headings)'  # Use a clean, modern font
        run.font.size = Pt(24)
        run.bold = True
        # Align to the right
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return self.doc
    
    def add_company_header(self, company_name: str) -> None:
        """Add company header to document"""
        if not self.doc:
            raise ValueError("Document not created. Call create_document() first.")
        
        self.doc.add_heading(f"{company_name}:", level=1)
    
    def add_company_info_section(self, company_info: CompanyInfo) -> Table:
        """Add company information section"""
        self.doc.add_heading("Company Information", level=2)
        
        data = [
            (f"Company Name", company_info.name),
            (f"Country of Domicile", company_info.country),
            (f"Analyst", company_info.analyst),
            (f"Update", company_info.update_date),
            (f"GICS Sector", company_info.gics_sector),
            (f"GICS Group", company_info.gics_industry_group),
            (f"Industry", company_info.gics_industry),
            (f"Sub Industry", company_info.gics_sub_industry)
        ]
        
        # Split into two columns
        _split_idx = int(round(len(data)/2,1))
        paired_data = [(data[i], data[i+_split_idx]) for i in range(_split_idx)]
        formatted_data = []
        for left, right in paired_data:
            formatted_data.append((f"{left[0]}: {left[1]}", f"{right[0]}: {right[1]}"))
        
        return self.table_builder.build_key_value_table(self.doc, formatted_data)
    
    def add_recommendation_section(self, recommendation: InvestmentRecommendation) -> Table:
        """Add investment recommendation section"""
        self.doc.add_heading("Internal Recommendation", level=2)
        
        data = [
            (f"Buy/Sell Rec: {recommendation.buy_sell_rec}", 
             f"Last Price: {recommendation.last_price}"),
            (f"Idea Stage: {recommendation.idea_stage}", 
             f"Base Target Price: {recommendation.base_target}"),
            (f"ESG Rating: {recommendation.esg_rating}", 
             f"Bull Target Price: {recommendation.bull_target}"),
            (f"Investment Theme: {recommendation.theme}", 
             f"Bear Target Price: {recommendation.bear_target}")
        ]
        
        return self.table_builder.build_key_value_table(self.doc, data)
    
    def add_thesis_section(self, thesis_title: str, thesis_text: str) -> Table:
        """Add investment thesis section"""
        self.doc.add_paragraph("")  # Add spacing
        
        data = [thesis_title, thesis_text]
        return self.table_builder.build_single_column_table(self.doc, data)
    
    def add_financials_section(self, title: str, financial_data: List[List[str]]) -> Table:
        """Add financials section"""
        self.doc.add_heading(title, level=2)
        return self.table_builder.build_data_table(self.doc, financial_data, has_header=True)
    
    def add_portfolio_exposure_section(self, title: str, exposure_data: List[List[str]]) -> Table:
        """Add financials section"""
        self.doc.add_heading(title, level=2)
        return self.table_builder.build_data_table(self.doc, exposure_data, has_header=True)
    
    def save_document(self, filename: str) -> None:
        """Save the document"""
        if not self.doc:
            raise ValueError("No document to save")
        
        self.doc.save(filename)




def convert_docx_to_pdf_silently(input_path: str, output_path: str):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False   # Don't show Word UI
    word.DisplayAlerts = 0  # Disable alerts

    doc = word.Documents.Open(os.path.abspath(input_path))
    doc.SaveAs(os.path.abspath(output_path), FileFormat=17)  # 17 = PDF
    doc.Close()
    word.Quit()


def create_thesis_doc():
    """Create a sample research report"""
    
    # Create custom styling
    custom_style = TableStyle(
        font=FontStyle(name='Calibri', size=10),
        alternate_shading=True,
        shade_color="F2F2F2",
        shade_header=True  # Enable header shading
    )
    
    # Initialize generator
    generator = BloombergReportGenerator(custom_style)
    generator.create_document()
    
    # Company information
    company_info = CompanyInfo(
        name="Apple Inc",
        country="US",
        analyst="Lucas Escobar",
        gics_sector="Information Technology",
        gics_industry_group="Technology, Hardware, & Equipment",
        gics_industry="Technology Hardware, Storage &",
        gics_sub_industry="Technology Hardware, Storage &"
    )
    
    # Investment recommendation
    recommendation = InvestmentRecommendation(
        buy_sell_rec="Buy",
        idea_stage="ACTIVE",
        esg_rating="Leading",
        theme="Tech Trifecta",
        last_price=203.92,
        base_target=300.00,
        bull_target=355.00,
        bear_target=200.00
    )
    
    # Build report
    generator.add_company_header("Apple Inc")
    generator.add_company_info_section(company_info)
    generator.add_recommendation_section(recommendation)
    
    thesis_text = """Apple's near-term growth could land in the low-single digits, driven by a steady services 
business that thrives on an installed base of 2.2 billion active devices. This is even as product sales remain depressed. 
Consumer weakness in China and increased competition are hampering near-term sales."""
    
    generator.add_thesis_section("Investment Thesis:", thesis_text)
    
    # Financial data
    financials_data = [
        ["", "2022 A", "2023 A", "2024 A*", "2025 A", "2026 A", "2027 A", "Est Source"],
        ["Sales", "394328", "383285", "391035", "94832", "100113", "453388", "Internal (EQ)"],
        ["EPS Adj", "6.11", "6.13", "6.75", "7.16", "7.64", "8.44", "Consensus"],
        ["EPS GAAP", "6.11", "6.13", "6.08", "6", "7", "8.44", "Internal (EQ)"],
        ["Gross Profit", "170782", "169148", "180683", "---", "---", "---", ""],
        ["Gross Margin", "43.31", "44.13", "46.21", "46.55", "46.55", "47.15", "Consensus"],
        ["EBITDA", "130541", "125820", "134661", "0", "0", "158062", "Internal (EQ)"],
        ["EBIT", "119437", "114301", "123216", "127791", "132856", "143773", "Consensus"],
        ["Net Income, Adj", "99803", "96995", "103998", "107406", "111615", "121776", "Consensus"],
        ["Net Income, GAAP", "99803", "96995", "93736", "9843", "11094", "121776", "Internal (EQ)"],
        ["CAPEX", "-10708", "-10595", "-9447", "-11053", "-12261", "-11400", "Consensus"],
        ["FCF", "111443", "99584", "108807", "104200", "121756", "128150", "Consensus"],
    ]
    
    generator.add_financials_section("Financials & Forecasts", financials_data)
    
    portfolio_exposure_data = [
        ["ID", "Port Wgt", "Bench Wgt", "Active Wgt", "TotRtn Port", "TotRtn Bench", "TotRtn Active", "Tot Attr"],
        ["EQUITY8_US", "5.45", "5.21", "-0.24", "1.53", "1.53", "0", "0"],
        ["EQUITY8_US_VALUE", "0.07", "0", "-0.07", "1.53", "0", "1.53", "0"],
        ["EQUITY8_CANADIAN", "4.06", "0", "-4.06", "1.53", "0", "1.53", "0"],
        ["EQUITY8_EM", "0", "0", "0", "0", "0", "0", "0"],
        ["EQUITY8_ESG", "5.15", "5.76", "0.61", "1.53", "1.53", "0", "0"],
        ["EQUITY8_GLOBAL", "3.88", "4.18", "0.31", "1.53", "1.53", "0", "0"],
        ["EQUITY8_LONG_SHORT", "4.94", "5.45", "0.51", "1.53", "1.53", "0", "0"],
        ["EQUITY8_MID_CAP_GROWTH", "0.1", "0", "-0.1", "1.53", "0", "1.53", "0"],
        ["EQUITY8_SMALL_CAP_GROWTH", "0", "0", "0", "0", "0", "0", "0"]
    ]

    generator.add_portfolio_exposure_section("Portfolio Exposure", portfolio_exposure_data)

    generator.save_document("thesis_doc.docx")
    convert_docx_to_pdf_silently("thesis_doc.docx", "thesis_doc.pdf")


