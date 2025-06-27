"""
research_creation_system/document_generation/text_formatter.py

Handles advanced text formatting operations for professional document layout.
Provides clean separation between table building and text formatting concerns.
"""

from typing import List, Dict, Any, Optional, Tuple, Union
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .styling import TableStyle, AlignmentType, FontStyle
from .document_styler import DocumentStyler

class TextFormatter:
    """Handles sophisticated text formatting for professional documents"""


    def __init__(self, styler: DocumentStyler):
        self.styler = styler

    def add_formatted_paragraph(self, doc: Document, text: str, bold: bool = False,
                            italic: bool = False, size: int = None,
                            spacing_before: float = 0, spacing_after: float = 0) -> None:
        """Add a formatted paragraph with custom styling"""
        paragraph = doc.add_paragraph()
    
        # Set paragraph spacing
        if spacing_before > 0:
            paragraph.paragraph_format.space_before = Pt(spacing_before)
        if spacing_after > 0:
            paragraph.paragraph_format.space_after = Pt(spacing_after)
        else:
            # Minimal default spacing for tighter layout
            paragraph.paragraph_format.space_after = Pt(3)
    
        # Add formatted text
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        else:
            run.font.size = Pt(self.styler.style_config.font.size)
        run.font.name = self.styler.style_config.font.name

    def add_bulleted_list(self, doc: Document, items: List[str], indent_level: int = 0,
                        bold_items: bool = False, italic_items: bool = False,
                        compact: bool = True) -> None:
        """Add a properly formatted bulleted list"""
        for item in items:
            paragraph = doc.add_paragraph()
        
            # Set tight spacing for compact lists
            if compact:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
            
            # Set indentation
            paragraph.paragraph_format.left_indent = Inches(0.25 + (indent_level * 0.25))
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        
            # Add bullet and text
            run = paragraph.add_run(f"• {item}")
            run.bold = bold_items
            run.italic = italic_items
            run.font.name = self.styler.style_config.font.name
            run.font.size = Pt(self.styler.style_config.font.size)

    def add_mixed_format_paragraph(self, doc: Document,
                                segments: List[Tuple[str, bool, bool]]) -> None:
        """Add paragraph with mixed formatting segments (text, bold, italic)"""
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)  # Tight spacing
    
        for text, bold, italic in segments:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.name = self.styler.style_config.font.name
            run.font.size = Pt(self.styler.style_config.font.size)

    def add_subsection_with_bullets(self, doc: Document, title: str, items: List[str],
                                title_bold: bool = True, items_italic: bool = False) -> None:
        """Add a subsection with title and bulleted items"""
        # Add title
        self.add_formatted_paragraph(doc, title, bold=title_bold, spacing_after=2)
    
        # Add bulleted items with minimal spacing
        self.add_bulleted_list(doc, items, italic_items=items_italic, compact=True)