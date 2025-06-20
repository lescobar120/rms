# TableBuilder class - builds different types of tables
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass, field
from enum import Enum
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from .document_styler import DocumentStyler
from .styling import AlignmentType, FontStyle



class TableBuilder:
    """Builds different types of tables with consistent styling"""
    
    def __init__(self, styler: DocumentStyler):
        self.styler = styler
    
    # def build_key_value_table(self, doc: Document, data: List[Tuple[str, str]]) -> Table:
    #     """Build a two-column key-value table"""
    #     table = doc.add_table(rows=len(data), cols=2)
        
    #     for row_idx, (key, value) in enumerate(data):
    #         # Key column
    #         key_cell = table.cell(row_idx, 0)
    #         self.styler.style_cell_text(key_cell, f"{key}", AlignmentType.LEFT)
            
    #         # Value column (italic)
    #         value_cell = table.cell(row_idx, 1)
    #         value_font = FontStyle(italic=True)
    #         self.styler.style_cell_text(value_cell, value, AlignmentType.LEFT, value_font)
        
    #     self.styler.style_table(table, "key_value")
    #     self._set_column_widths(table, [Inches(3.75), Inches(3.75)])
    #     return table
    
    def build_key_value_table(self, doc: Document, data: List[Tuple[Tuple[str, str], Tuple[str, str]]]) -> Table:
        """
        Build a two-column key-value table with mixed formatting
        
        Args:
            data: List of ((left_key, left_value), (right_key, right_value)) tuples
        """
        table = doc.add_table(rows=len(data), cols=2)
        
        for row_idx, (left_pair, right_pair) in enumerate(data):
            left_key, left_value = left_pair
            right_key, right_value = right_pair
            
            # Left column - mixed formatting
            left_cell = table.cell(row_idx, 0)
            self._style_key_value_cell(left_cell, left_key, left_value)
            
            # Right column - mixed formatting
            right_cell = table.cell(row_idx, 1)
            self._style_key_value_cell(right_cell, right_key, right_value)
        
        self.styler.style_table(table, "key_value")
        return table

    def _style_key_value_cell(self, cell, key: str, value: str):
        """Style a cell with key (regular) and value (italic)"""
        p = cell.paragraphs[0]
        p.alignment = AlignmentType.LEFT.value
        p.clear()
        
        # Add key part (regular)
        key_run = p.add_run(f"{key}: ")
        key_run.font.name = self.styler.style_config.font.name
        key_run.font.size = Pt(self.styler.style_config.font.size)
        key_run.bold = False
        key_run.italic = False
        
        # Add value part (italic)
        value_run = p.add_run(value)
        value_run.font.name = self.styler.style_config.font.name
        value_run.font.size = Pt(self.styler.style_config.font.size) 
        value_run.bold = False
        value_run.italic = True
    
    def build_data_table(self, doc: Document, data: List[List[str]], 
                        has_header: bool = True) -> Table:
        """Build a multi-column data table"""
        if not data:
            raise ValueError("Data cannot be empty")
        
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        
        for row_idx, row_data in enumerate(data):
            is_header = has_header and row_idx == 0
            
            for col_idx, cell_value in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                
                # Determine alignment
                if is_header:
                    if col_idx == 0:
                        alignment = AlignmentType.LEFT
                    else:
                        alignment = AlignmentType.CENTER
                    font_style = FontStyle(bold=True)
                elif col_idx == 0:
                    alignment = AlignmentType.LEFT
                    font_style = FontStyle()
                else:
                    alignment = AlignmentType.RIGHT
                    font_style = FontStyle()
                
                self.styler.style_cell_text(cell, str(cell_value), alignment, font_style)
        
        self.styler.style_table(table, "data")
        return table

    def build_single_column_table(self, doc: Document, data: List[str]) -> Table:
        """Build a single-column table (useful for thesis/description blocks)"""
        table = doc.add_table(rows=len(data), cols=1)
        
        for row_idx, text in enumerate(data):
            cell = table.cell(row_idx, 0)
            font_style = FontStyle(bold=(row_idx == 0))  # First row bold
            self.styler.style_cell_text(cell, text, AlignmentType.LEFT, font_style)
        
        self.styler.style_table(table, "thesis")
        return table
    
    def _set_column_widths(self, table: Table, widths: List[Inches]) -> None:
        """Set column widths for a table"""
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(widths):
                    cell.width = widths[col_idx]

