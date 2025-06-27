from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass, field
from enum import Enum
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .styling import TableStyle
from .document_styler import DocumentStyler
from .table_builder import TableBuilder
from .text_formatter import TextFormatter
from ..data_models import CompanyInfo, InvestmentRecommendation



class BloombergReportGenerator:
    """Main class for generating Bloomberg RMS NOTEs"""
    
    def __init__(self, style_config: TableStyle = None):
        self.styler = DocumentStyler(style_config)
        self.table_builder = TableBuilder(self.styler)
        self.text_formatter = TextFormatter(self.styler)
        self.doc = None
    
    def create_document(self, margins: float = 0.5) -> Document:
        """Create a new document with custom margins"""
        self.doc = Document()
        
        # Set margins
        section = self.doc.sections[0]
        section.top_margin = Inches(margins)
        section.bottom_margin = Inches(margins)
        section.left_margin = Inches(margins)
        section.right_margin = Inches(margins)

        # Add a header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = header_para.add_run("Bloomberg")
        run.font.name = 'Calibri (Headings)'  # Use a clean, modern font
        run.font.size = Pt(24)
        run.bold = True
        # Align to the right
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return self.doc
    
    def add_company_header(self, company_name: str) -> None:
        """Add company header to document"""
        if not self.doc:
            raise ValueError("Document not created. Call create_document() first.")
        
        self.doc.add_heading(f"{company_name}:", level=1)
    
    # def add_company_info_section(self, company_info: CompanyInfo) -> Table:
    #     """Add company information section"""
    #     self.doc.add_heading("Company Information", level=2)
        
    #     data = [
    #         (f"Company Name", company_info.name),
    #         (f"Country of Domicile", company_info.country),
    #         (f"Analyst", company_info.analyst),
    #         (f"Update", company_info.update_date),
    #         (f"GICS Sector", company_info.gics_sector),
    #         (f"GICS Group", company_info.gics_industry_group),
    #         (f"Industry", company_info.gics_industry),
    #         (f"Sub Industry", company_info.gics_sub_industry)
    #     ]
        
    #     # Split into two columns
    #     _split_idx = int(round(len(data)/2,1))
    #     paired_data = [(data[i], data[i+_split_idx]) for i in range(_split_idx)]
    #     formatted_data = []
    #     for left, right in paired_data:
    #         formatted_data.append((f"{left[0]}: {left[1]}", f"{right[0]}: {right[1]}"))
        
    #     return self.table_builder.build_key_value_table(self.doc, formatted_data)
    
    def add_company_info_section(self, company_info: CompanyInfo) -> Table:
        """Add company information section"""
        self.doc.add_heading("Company Information", level=2)
        
        # Create structured data as tuples of (key, value) pairs
        data = [
            ("Company Name", company_info.name),
            ("Country of Domicile", company_info.country),
            ("Analyst", company_info.analyst),
            ("Update", company_info.update_date),
            ("GICS Sector", company_info.gics_sector),
            ("GICS Group", company_info.gics_industry_group),
            ("Industry", company_info.gics_industry),
            ("Sub Industry", company_info.gics_sub_industry)
        ]
        
        # Split into two columns of (key, value) pairs
        split_idx = int(round(len(data)/2, 1))
        paired_data = []
        
        for i in range(split_idx):
            left_pair = data[i]  # (key, value)
            right_pair = data[i + split_idx] if i + split_idx < len(data) else ("", "")
            paired_data.append((left_pair, right_pair))
        
        return self.table_builder.build_key_value_table(self.doc, paired_data)
    
    # def add_recommendation_section(self, recommendation: InvestmentRecommendation) -> Table:
    #     """Add investment recommendation section"""
    #     self.doc.add_heading("Internal Recommendation", level=2)
        
    #     data = [
    #         (f"Buy/Sell Rec: {recommendation.buy_sell_rec}", 
    #          f"Last Price: {recommendation.last_price}"),
    #         (f"Idea Stage: {recommendation.idea_stage}", 
    #          f"Base Target Price: {recommendation.base_target}"),
    #         (f"ESG Rating: {recommendation.esg_rating}", 
    #          f"Bull Target Price: {recommendation.bull_target}"),
    #         (f"Investment Theme: {recommendation.theme}", 
    #          f"Bear Target Price: {recommendation.bear_target}")
    #     ]
        
    #     return self.table_builder.build_key_value_table(self.doc, data)

    def add_recommendation_section(self, recommendation: InvestmentRecommendation) -> Table:
        """Add investment recommendation section"""
        self.doc.add_heading("Internal Recommendation", level=2)
        
        # Create structured data as tuples of (key, value) pairs
        data = [
            ("Buy/Sell Rec", recommendation.buy_sell_rec),
            ("Last Price", str(recommendation.last_price)),
            ("Idea Stage", recommendation.idea_stage), 
            ("Base Target Price", str(recommendation.base_target)),
            ("ESG Rating", recommendation.esg_rating),
            ("Bull Target Price", str(recommendation.bull_target)),
            ("Investment Theme", recommendation.theme),
            ("Bear Target Price", str(recommendation.bear_target))
        ]
        
        # Split into two columns of (key, value) pairs
        split_idx = int(round(len(data)/2, 1))
        paired_data = []
        
        for i in range(split_idx):
            left_pair = data[i]  # (key, value)
            right_pair = data[i + split_idx] if i + split_idx < len(data) else ("", "")
            paired_data.append((left_pair, right_pair))
        
        # Use the new mixed format method
        return self.table_builder.build_key_value_table(self.doc, paired_data)
    
    def add_thesis_section(self, thesis_title: str, thesis_text: str) -> Table:
        """Add investment thesis section"""
        self.doc.add_paragraph("")  # Add spacing
        
        data = [thesis_title, thesis_text]
        return self.table_builder.build_single_column_table(self.doc, data)
    
    def add_financials_section(self, title: str, financial_data: List[List[str]]) -> Table:
        """Add financials section"""
        self.doc.add_heading(title, level=2)
        return self.table_builder.build_data_table(self.doc, financial_data, has_header=True)
    
    def add_portfolio_exposure_section(self, title: str, exposure_data: List[List[str]]) -> Table:
        """Add financials section"""
        self.doc.add_heading(title, level=2)
        return self.table_builder.build_data_table(self.doc, exposure_data, has_header=True)
    

    def add_earnings_preview_section(self, title: str, earnings_date: str) -> None:
        """Add earnings preview title section"""
        self.doc.add_heading(title, level=1)
        self.doc.add_paragraph(f"**Earnings Date:** {earnings_date}")

    def add_estimates_comparison_section(self, estimates_data: List[List[str]]) -> Table:
        """Add estimates vs consensus comparison table"""
        self.doc.add_heading("Our Estimates vs. Consensus", level=2)
        return self.table_builder.build_data_table(self.doc, estimates_data, has_header=True)

    # def add_scenario_analysis_section(self, scenarios_data: List[List[str]], scenario_details: List[str] = None) -> Table:
    #     """Add scenario analysis table"""
    #     self.doc.add_heading("Scenario Analysis", level=2)
    #     table = self.table_builder.build_data_table(self.doc, scenarios_data, has_header=True)

    #     # Add scenario details as paragraphs if provided
    #     if scenario_details:
    #         for detail in scenario_details:
    #             self.doc.add_paragraph(detail)

    #     return table

    def add_scenario_analysis_section(self, scenarios_data: List[List[str]],
                                    scenario_details: List[str] = None) -> Table:
        """Add scenario analysis table with improved formatting"""
        self.doc.add_heading("Scenario Analysis", level=2)
        table = self.table_builder.build_data_table(self.doc, scenarios_data, has_header=True)
    
        # Add scenario details with better formatting
        if scenario_details:
            for detail in scenario_details:
                # Parse and format with mixed bold/italic
                if ":**" in detail:
                    parts = detail.split(":**", 1)
                    title = parts[0].replace("**", "") + ":"
                    assumptions = parts[1].strip()
                
                    segments = [
                        (title, True, False),     # Bold title
                        (" ", False, False),      # Space
                        (assumptions, False, True) # Italic assumptions
                    ]
                    self.text_formatter.add_mixed_format_paragraph(self.doc, segments)
                else:
                    self.text_formatter.add_formatted_paragraph(self.doc, detail, spacing_after=2)
    
        return table


    # def add_key_metrics_section(self, metrics: List[str]) -> None:
    #     """Add key metrics to watch section"""
    #     self.doc.add_heading("Key Metrics to Watch", level=2)
    #     for metric in metrics:
    #         self.doc.add_paragraph(f"• **{metric}**",)

    def add_key_metrics_section(self, metrics: List[str]) -> None:
        """Add key metrics section with improved formatting"""
        self.doc.add_heading("Key Metrics to Watch", level=2)
        self.text_formatter.add_bulleted_list(self.doc, metrics, bold_items=True, compact=True)

    def add_historical_earnings_section(self, historical_data: List[List[str]]) -> Table:
        """Add historical earnings performance table"""
        self.doc.add_heading("Historical Earnings Performance", level=2)
        return self.table_builder.build_data_table(self.doc, historical_data, has_header=True)

    # def add_management_focus_section(self, guidance_items: List[str] = None, questions: List[str] = None) -> None:
    #     """Add management guidance and Q&A focus section"""
    #     self.doc.add_heading("Management Guidance & Key Questions", level=2)

    #     if guidance_items:
    #         self.doc.add_paragraph("**Expected Guidance Updates:**")
    #         for item in guidance_items:
    #             self.doc.add_paragraph(f"• {item}")

    #     if questions:
    #         if guidance_items:
    #             self.doc.add_paragraph("")  # Add spacing
    #         self.doc.add_paragraph("**Key Questions for Management:**")
    #         for question in questions:
    #             self.doc.add_paragraph(f"• {question}")

    
    def add_management_focus_section(self, guidance_items: List[str] = None,
                                questions: List[str] = None) -> None:
        """Add management focus section with improved formatting"""
        self.doc.add_heading("Management Guidance & Key Questions", level=2)
    
        if guidance_items:
            self.text_formatter.add_subsection_with_bullets(
                self.doc, "Expected Guidance Updates:", guidance_items,
                title_bold=True, items_italic=False
            )
    
        if questions:
            if guidance_items:
                self.text_formatter.add_formatted_paragraph(self.doc, "", spacing_after=6)
        
            self.text_formatter.add_subsection_with_bullets(
                self.doc, "Key Questions for Management:", questions,
                title_bold=True, items_italic=False
            )


    def add_risks_catalysts_section(self, risks: List[str] = None, catalysts: List[str] = None) -> None:
        """Add risks and catalysts section"""
        self.doc.add_heading("Risks & Catalysts", level=2)

        if risks:
            self.doc.add_paragraph("**Key Risks:**")
            for risk in risks:
                self.doc.add_paragraph(f"• {risk}")

        if catalysts:
            if risks:
                self.doc.add_paragraph("")  # Add spacing
            self.doc.add_paragraph("**Potential Catalysts:**")
            for catalyst in catalysts:
                self.doc.add_paragraph(f"• {catalyst}")

    def add_risks_catalysts_section(self, risks: List[str] = None,
                                catalysts: List[str] = None) -> None:
        """Add management focus section with improved formatting"""
        self.doc.add_heading("Management Guidance & Key Questions", level=2)
    
        if risks:
            self.text_formatter.add_subsection_with_bullets(
                self.doc, "Key Risks:", risks,
                title_bold=True, items_italic=False
            )
    
        if catalysts:
            if risks:
                self.text_formatter.add_formatted_paragraph(self.doc, "", spacing_after=6)
        
            self.text_formatter.add_subsection_with_bullets(
                self.doc, "Potential Catalysts:", catalysts,
                title_bold=True, items_italic=False
            )


    def add_market_context_section(self, sector_context: str = None, competitive_context: str = None) -> None:
        """Add market and competitive context section"""
        self.doc.add_heading("Market Context", level=2)


        # if sector_context:
        #     self.doc.add_paragraph(f"**Sector Context:** {sector_context}")

        # if competitive_context:
        #     if sector_context:
        #         self.doc.add_paragraph("")  # Add spacing
        #     self.doc.add_paragraph(f"**Competitive Positioning:** {competitive_context}")

        
        if sector_context:
            sector_context = f"**Sector Context:** {sector_context}"
            # Parse and format with mixed bold/italic
            parts = sector_context.split(":**", 1)
            title = parts[0].replace("**", "") + ":"
            assumptions = parts[1].strip()
        
            segments = [
                (title, True, False),     # Bold title
                (" ", False, False),      # Space
                (assumptions, False, True) # Italic assumptions
            ]
            self.text_formatter.add_mixed_format_paragraph(self.doc, segments)

        if competitive_context:
            competitive_context = f"**Competitive Positioning:** {competitive_context}"
            # Parse and format with mixed bold/italic
            parts = competitive_context.split(":**", 1)
            title = parts[0].replace("**", "") + ":"
            assumptions = parts[1].strip()
        
            segments = [
                (title, True, False),     # Bold title
                (" ", False, False),      # Space
                (assumptions, False, True) # Italic assumptions
            ]
            self.text_formatter.add_mixed_format_paragraph(self.doc, segments)

    
    def save_document(self, filename: str) -> None:
        """Save the document"""
        if not self.doc:
            raise ValueError("No document to save")
        
        self.doc.save(filename)




